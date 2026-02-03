from fastapi import APIRouter, Depends, UploadFile, File
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session as OrmSession
from sqlalchemy import asc
import os
import subprocess
import tempfile
import re
from pathlib import Path
from datetime import datetime

from ..database import get_db
from ..models import Session as SessionModel, Content as ContentModel
from ..utils import ok, err, dt_str

router = APIRouter(prefix="/content", tags=["content"])


@router.get("/get/{session_id}")
def get_contents(session_id: int, db: OrmSession = Depends(get_db)):
    s = db.get(SessionModel, session_id)
    if not s:
        return err(404, "会话不存在")

    rows = (
        db.query(ContentModel)
        .filter(ContentModel.session_id == session_id)
        .order_by(asc(ContentModel.created_at))
        .all()
    )
    data = [
        {
            "content_id": r.content_id,
            "content": r.content,
            "content_type": r.content_type,
            "created_at": dt_str(r.created_at),
        }
        for r in rows
    ]
    return ok(data)


@router.delete("/clear/{session_id}")
def clear_session_content(session_id: int, db: OrmSession = Depends(get_db)):
    """
    清空会话的所有内容
    """
    try:
        # 检查会话是否存在
        s = db.get(SessionModel, session_id)
        if not s:
            return err(404, "会话不存在")
        
        # 删除该会话的所有内容
        deleted_count = db.query(ContentModel).filter(ContentModel.session_id == session_id).delete()
        db.commit()
        
        print(f"[DEBUG] 清空会话 {session_id} 的内容，删除了 {deleted_count} 条记录")
        
        return ok(None, f"已清空 {deleted_count} 条内容")
    
    except Exception as e:
        print(f"[DEBUG] 清空内容失败: {str(e)}")
        return err(500, f"清空失败: {str(e)}")


@router.post("/upload")
async def upload_reference_document(file: UploadFile = File(...)):
    """
    上传参考文档，支持.docx和.md文件
    .docx文件会使用pandoc转换为markdown
    """
    try:
        # 检查文件类型
        file_extension = Path(file.filename).suffix.lower()
        
        if file_extension not in ['.docx', '.md']:
            return err(400, "不支持的文件类型，请上传.docx或.md文件")
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            # 保存上传的文件
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        # 如果是.docx文件，转换为markdown
        if file_extension == '.docx':
            try:
                # 使用pandoc转换.docx为markdown
                output_md_path = temp_file_path.replace('.docx', '.md')
                result = subprocess.run(
                    ['pandoc', '-f', 'docx', '-t', 'markdown', '-o', output_md_path, temp_file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                
                if result.returncode != 0:
                    print(f"[DEBUG] Pandoc转换失败: {result.stderr}")
                    return err(500, f"文档转换失败: {result.stderr}")
                
                # 读取转换后的markdown内容
                with open(output_md_path, 'r', encoding='utf-8') as f:
                    markdown_content = f.read()
                
                # 清理临时文件
                os.unlink(temp_file_path)
                os.unlink(output_md_path)
                
                return ok({
                    "filename": file.filename,
                    "content": markdown_content,
                    "type": "markdown"
                }, "文档上传并转换成功")
                
            except subprocess.TimeoutExpired:
                return err(500, "文档转换超时")
            except Exception as e:
                print(f"[DEBUG] 转换过程中出错: {str(e)}")
                return err(500, f"文档转换失败: {str(e)}")
        
        # 如果是.md文件，直接读取内容
        elif file_extension == '.md':
            try:
                with open(temp_file_path, 'r', encoding='utf-8') as f:
                    markdown_content = f.read()
                
                # 清理临时文件
                os.unlink(temp_file_path)
                
                return ok({
                    "filename": file.filename,
                    "content": markdown_content,
                    "type": "markdown"
                }, "文档上传成功")
                
            except Exception as e:
                print(f"[DEBUG] 读取markdown文件出错: {str(e)}")
                return err(500, f"文档读取失败: {str(e)}")
        
    except Exception as e:
        print(f"[DEBUG] 文件上传过程中出错: {str(e)}")
        return err(500, f"文件上传失败: {str(e)}")


@router.get("/export/{session_id}")
def export_document(session_id: int, export_type: str = "md", reference_doc: str = None, db: OrmSession = Depends(get_db)):
    """
    导出会话内容为文档
    export_type: "md" 或 "docx"
    reference_doc: 参考文档的文件名（如果有docx参考文档，使用它作为模板）
    """
    try:
        # 验证导出类型
        if export_type not in ['md', 'docx']:
            return err(400, "不支持的导出类型，请选择 md 或 docx")
        
        # 获取会话信息
        s = db.get(SessionModel, session_id)
        if not s:
            return err(404, "会话不存在")
        
        # 获取会话的所有内容
        rows = (
            db.query(ContentModel)
            .filter(ContentModel.session_id == session_id)
            .order_by(asc(ContentModel.created_at))
            .all()
        )
        
        # 只获取最新的内容（AI生成的回复）
        if not rows:
            return err(404, "该会话没有可导出的内容")
        
        # 只导出最新的内容，去除可能的前缀（如"标题："、"正文："等）
        latest_content = rows[-1].content
        
        # 清理内容：去除"标题："、"正文："等前缀
        cleaned_content = re.sub(r'^(标题：|正文：|Title:|Content:)\s*', '', latest_content, flags=re.MULTILINE)
        
        # 去除开头和结尾的空白
        cleaned_content = cleaned_content.strip()
        
        # 创建临时文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if export_type == "md":
            # 导出为markdown文件
            temp_file = tempfile.NamedTemporaryFile(
                mode='w',
                delete=False,
                suffix='.md',
                encoding='utf-8'
            )
            temp_file.write(cleaned_content)
            temp_file.close()
            temp_file_path = temp_file.name
            
            filename = f"{s.session_name}_{timestamp}.md"
            
            return FileResponse(
                path=temp_file_path,
                media_type='text/markdown',
                filename=filename
            )
        
        elif export_type == "docx":
            # 导出为docx文件
            # 如果有参考文档，直接将内容写入参考文档中
            if reference_doc:
                # 获取项目根目录（从当前文件向上查找）
                current_dir = os.path.dirname(os.path.abspath(__file__))
                # 向上查找项目根目录（找到format目录）
                while current_dir and not os.path.exists(os.path.join(current_dir, 'format')):
                    current_dir = os.path.dirname(current_dir)
                    if current_dir == os.path.dirname(current_dir):  # 到达根目录
                        break
                
                if current_dir:
                    reference_path = os.path.join(current_dir, 'format', reference_doc)
                    if os.path.exists(reference_path):
                        # 使用python-docx库修改参考文档
                        try:
                            from docx import Document
                            
                            # 打开参考文档
                            doc = Document(reference_path)
                            
                            # 清空文档内容
                            for paragraph in doc.paragraphs:
                                paragraph.text = ""
                            
                            # 添加新内容
                            doc.add_paragraph(cleaned_content)
                            
                            # 保存到临时文件
                            temp_docx_path = temp_md_path.replace('.md', '.docx')
                            doc.save(temp_docx_path)
                            
                            print(f"[DEBUG] 使用参考文档并写入内容: {reference_path}")
                            
                            filename = f"{s.session_name}_{timestamp}.docx"
                            
                            return FileResponse(
                                path=temp_docx_path,
                                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                filename=filename
                            )
                        except ImportError:
                            print(f"[DEBUG] python-docx库未安装，使用pandoc转换")
                        except Exception as e:
                            print(f"[DEBUG] 使用python-docx失败: {str(e)}")
            
            # 如果没有参考文档或python-docx不可用，使用pandoc转换
            # 先创建临时markdown文件
            temp_md_file = tempfile.NamedTemporaryFile(
                mode='w',
                delete=False,
                suffix='.md',
                encoding='utf-8'
            )
            temp_md_file.write(cleaned_content)
            temp_md_file.close()
            temp_md_path = temp_md_file.name
            
            # 使用pandoc转换为docx
            temp_docx_path = temp_md_path.replace('.md', '.docx')
            result = subprocess.run(
                ['pandoc', '-f', 'markdown', '-t', 'docx', '-o', temp_docx_path, temp_md_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            # 清理临时markdown文件
            os.unlink(temp_md_path)
            
            if result.returncode != 0:
                print(f"[DEBUG] Pandoc转换失败: {result.stderr}")
                return err(500, f"文档转换失败: {result.stderr}")
            
            filename = f"{s.session_name}_{timestamp}.docx"
            
            return FileResponse(
                path=temp_docx_path,
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                filename=filename
            )
    
    except Exception as e:
        print(f"[DEBUG] 导出文档过程中出错: {str(e)}")
        return err(500, f"导出文档失败: {str(e)}")
